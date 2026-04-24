import os
import logging
import hashlib
import tempfile
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
import asyncio
import aiofiles
import PyPDF2
import docx
try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
from pathlib import Path
import time

from models.schemas import DocumentInfo, DocumentType, DocumentStatus
from utils.database import get_db_connection

logger = logging.getLogger(__name__)

class DocumentService:
    def __init__(self):
        self.chunk_size = int(os.getenv("CHUNK_SIZE", "1000"))
        self.chunk_overlap = int(os.getenv("CHUNK_OVERLAP", "200"))
        self.max_file_size = int(os.getenv("MAX_FILE_SIZE_MB", "50")) * 1024 * 1024
        self.upload_dir = os.getenv("UPLOAD_DIR", "./data/uploads")
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        # Ensure upload directory exists
        os.makedirs(self.upload_dir, exist_ok=True)
    
    async def process_document(
        self, 
        file_content: bytes, 
        filename: str,
        vector_service,
        embedding_service,
        llm_service
    ) -> Dict[str, Any]:
        """Process uploaded document end-to-end"""
        start_time = time.time()
        
        try:
            # Validate file
            await self._validate_file(file_content, filename)
            
            # Generate document ID
            document_id = self._generate_document_id(filename, file_content)
            
            # Determine file type
            file_type = self._get_file_type(filename)
            
            # Extract text from file
            text_content = await self._extract_text(file_content, filename, file_type)
            
            if not text_content.strip():
                raise ValueError("No text content found in the document")
            
            # Generate title and summary
            title = await llm_service.generate_title(text_content[:1000])
            
            # Split text into chunks
            chunks = self._split_text(text_content)
            
            if not chunks:
                raise ValueError("Failed to create text chunks")
            
            # Generate embeddings
            logger.info(f"Generating embeddings for {len(chunks)} chunks")
            embeddings = await embedding_service.encode_batch(chunks)
            
            if len(embeddings) != len(chunks):
                raise ValueError("Mismatch between chunks and embeddings")
            
            # Prepare metadata for each chunk
            metadata = []
            for i, chunk in enumerate(chunks):
                meta = {
                    "document_id": document_id,
                    "filename": filename,
                    "chunk_index": i,
                    "chunk_size": len(chunk),
                    "file_type": file_type.value,
                    "created_at": datetime.now().isoformat()
                }
                metadata.append(meta)
            
            # Store in vector database
            success = await vector_service.add_documents(
                document_id=document_id,
                chunks=chunks,
                embeddings=embeddings,
                metadata=metadata
            )
            
            if not success:
                raise ValueError("Failed to store document in vector database")
            
            # Save document info to database
            doc_info = DocumentInfo(
                id=document_id,
                filename=filename,
                file_size=len(file_content),
                content_preview=text_content[:500] + ("..." if len(text_content) > 500 else ""),
                chunk_count=len(chunks),
                type=file_type,
                status=DocumentStatus.COMPLETED,
                created_at=datetime.now(),
                updated_at=datetime.now(),
                metadata={
                    "title": title,
                    "text_length": len(text_content),
                    "processing_time": time.time() - start_time
                }
            )
            
            await self._save_document_info(doc_info)
            
            # Optionally save original file
            if os.getenv("SAVE_ORIGINAL_FILES", "false").lower() == "true":
                await self._save_original_file(file_content, document_id, filename)
            
            processing_time = time.time() - start_time
            logger.info(f"Document processed successfully in {processing_time:.2f}s")
            
            return {
                "id": document_id,
                "name": filename,
                "size": f"{len(file_content) / (1024 * 1024):.2f} MB",
                "uploaded": "Just now",
                "status": DocumentStatus.COMPLETED,
                "chunks": len(chunks),
                "type": file_type,
                "processing_time": processing_time,
                "title": title
            }
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            # Update status to failed if document was created
            try:
                await self._update_document_status(document_id, DocumentStatus.FAILED)
            except:
                pass
            raise
    
    async def _validate_file(self, file_content: bytes, filename: str):
        """Validate uploaded file"""
        # Check file size
        if len(file_content) > self.max_file_size:
            raise ValueError(f"File size exceeds {self.max_file_size / (1024*1024):.0f}MB limit")
        
        # Check file extension
        allowed_extensions = ['.pdf', '.txt', '.docx']
        file_extension = Path(filename).suffix.lower()
        
        if file_extension not in allowed_extensions:
            raise ValueError(f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}")
        
        # Check if file is empty
        if len(file_content) == 0:
            raise ValueError("File is empty")
    
    def _generate_document_id(self, filename: str, content: bytes) -> str:
        """Generate unique document ID"""
        content_hash = hashlib.md5(content).hexdigest()
        timestamp = datetime.now().isoformat()
        combined = f"{filename}_{timestamp}_{content_hash}"
        return hashlib.sha256(combined.encode()).hexdigest()[:16]
    
    def _get_file_type(self, filename: str) -> DocumentType:
        """Determine file type from filename"""
        extension = Path(filename).suffix.lower()
        
        if extension == '.pdf':
            return DocumentType.PDF
        elif extension == '.docx':
            return DocumentType.DOCX
        elif extension == '.txt':
            return DocumentType.TXT
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    async def _extract_text(self, file_content: bytes, filename: str, file_type: DocumentType) -> str:
        """Extract text from file based on type"""
        try:
            if file_type == DocumentType.PDF:
                return await self._extract_pdf_text(file_content)
            elif file_type == DocumentType.DOCX:
                return await self._extract_docx_text(file_content)
            elif file_type == DocumentType.TXT:
                return await self._extract_txt_text(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Text extraction failed for {filename}: {e}")
            raise ValueError(f"Failed to extract text from {file_type.value.upper()} file")
    
    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF"""
        text = ""
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
            tmp_file.write(file_content)
            tmp_file_path = tmp_file.name
        
        try:
            # Extract text using PyPDF2
            with open(tmp_file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue
        finally:
            # Clean up temporary file
            try:
                os.unlink(tmp_file_path)
            except:
                pass
        
        return text.strip()
    
    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX"""
        # Create temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            tmp_file.write(file_content)
            tmp_file_path = tmp_file.name
        
        try:
            # Extract text using python-docx
            doc = docx.Document(tmp_file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
            
        finally:
            # Clean up temporary file
            try:
                os.unlink(tmp_file_path)
            except:
                pass
    
    async def _extract_txt_text(self, file_content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            # Try UTF-8 first
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1
                return file_content.decode('latin-1')
            except UnicodeDecodeError:
                # Last resort - ignore errors
                return file_content.decode('utf-8', errors='ignore')
    
    def _split_text(self, text: str) -> List[str]:
        """Split text into chunks"""
        try:
            chunks = self.text_splitter.split_text(text)
            
            # Filter out very short chunks
            meaningful_chunks = [
                chunk.strip() for chunk in chunks 
                if len(chunk.strip()) >= 50  # Minimum chunk size
            ]
            
            return meaningful_chunks
            
        except Exception as e:
            logger.error(f"Text splitting failed: {e}")
            # Fallback to simple splitting
            return self._simple_text_split(text)
    
    def _simple_text_split(self, text: str) -> List[str]:
        """Simple fallback text splitting"""
        chunks = []
        words = text.split()
        
        current_chunk = []
        current_length = 0
        
        for word in words:
            word_length = len(word) + 1  # +1 for space
            
            if current_length + word_length > self.chunk_size and current_chunk:
                chunks.append(" ".join(current_chunk))
                current_chunk = [word]
                current_length = word_length
            else:
                current_chunk.append(word)
                current_length += word_length
        
        if current_chunk:
            chunks.append(" ".join(current_chunk))
        
        return [chunk for chunk in chunks if len(chunk.strip()) >= 50]
    
    async def _save_document_info(self, doc_info: DocumentInfo):
        """Save document information to database"""
        # This would typically save to a real database
        # For now, we'll use a simple file-based storage
        doc_file = os.path.join(self.upload_dir, f"{doc_info.id}_info.json")
        
        doc_data = {
            "id": doc_info.id,
            "filename": doc_info.filename,
            "file_size": doc_info.file_size,
            "content_preview": doc_info.content_preview,
            "chunk_count": doc_info.chunk_count,
            "type": doc_info.type.value,
            "status": doc_info.status.value,
            "created_at": doc_info.created_at.isoformat(),
            "updated_at": doc_info.updated_at.isoformat(),
            "metadata": doc_info.metadata
        }
        
        async with aiofiles.open(doc_file, 'w') as f:
            import json
            await f.write(json.dumps(doc_data, indent=2))
    
    async def _save_original_file(self, file_content: bytes, document_id: str, filename: str):
        """Save original file to disk"""
        file_path = os.path.join(self.upload_dir, f"{document_id}_{filename}")
        
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
    
    async def _update_document_status(self, document_id: str, status: DocumentStatus):
        """Update document status"""
        doc_file = os.path.join(self.upload_dir, f"{document_id}_info.json")
        
        if os.path.exists(doc_file):
            try:
                async with aiofiles.open(doc_file, 'r') as f:
                    import json
                    doc_data = json.loads(await f.read())
                
                doc_data['status'] = status.value
                doc_data['updated_at'] = datetime.now().isoformat()
                
                async with aiofiles.open(doc_file, 'w') as f:
                    await f.write(json.dumps(doc_data, indent=2))
                    
            except Exception as e:
                logger.error(f"Failed to update document status: {e}")
    
    async def get_all_documents(self) -> List[Dict[str, Any]]:
        """Get all processed documents"""
        documents = []
        
        try:
            for filename in os.listdir(self.upload_dir):
                if filename.endswith('_info.json'):
                    file_path = os.path.join(self.upload_dir, filename)
                    
                    try:
                        async with aiofiles.open(file_path, 'r') as f:
                            import json
                            doc_data = json.loads(await f.read())
                            documents.append(doc_data)
                    except Exception as e:
                        logger.warning(f"Failed to load document info from {filename}: {e}")
                        continue
            
            # Sort by creation date (newest first)
            documents.sort(key=lambda x: x.get('created_at', ''), reverse=True)
            
        except Exception as e:
            logger.error(f"Failed to get documents: {e}")
        
        return documents
    
    async def get_document_by_id(self, document_id: str) -> Optional[Dict[str, Any]]:
        """Get document by ID"""
        doc_file = os.path.join(self.upload_dir, f"{document_id}_info.json")
        
        if not os.path.exists(doc_file):
            return None
        
        try:
            async with aiofiles.open(doc_file, 'r') as f:
                import json
                return json.loads(await f.read())
        except Exception as e:
            logger.error(f"Failed to load document {document_id}: {e}")
            return None
    
    async def delete_document(self, document_id: str, vector_service) -> bool:
        """Delete document and its data"""
        try:
            # Delete from vector database
            await vector_service.delete_document(document_id)
            
            # Delete document info file
            doc_file = os.path.join(self.upload_dir, f"{document_id}_info.json")
            if os.path.exists(doc_file):
                os.unlink(doc_file)
            
            # Delete original file if it exists
            for filename in os.listdir(self.upload_dir):
                if filename.startswith(f"{document_id}_") and not filename.endswith('_info.json'):
                    file_path = os.path.join(self.upload_dir, filename)
                    os.unlink(file_path)
            
            logger.info(f"Document {document_id} deleted successfully")
            return True
            
        except Exception as e:
            logger.error(f"Failed to delete document {document_id}: {e}")
            return False